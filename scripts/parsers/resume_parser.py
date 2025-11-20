"""
Resume Parser - Extract text from PDF/DOCX files
Supports OCR for scanned PDFs and native text extraction
"""
import io
import logging
from typing import Dict, Optional
from datetime import datetime

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logging.warning("OCR dependencies not available. Install pytesseract and pdf2image for OCR support.")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx not available. DOCX parsing disabled.")

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logging.warning("PyPDF2 not available. PDF parsing disabled.")

try:
    from minio import Minio
    HAS_MINIO = True
except ImportError:
    HAS_MINIO = False
    logging.warning("MinIO client not available.")

import os
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Parser for extracting text from PDF and DOCX resume files.
    Supports both native text extraction and OCR for scanned documents.
    """
    
    def __init__(self, minio_endpoint: str = "minio:9000"):
        """
        Initialize parser with MinIO connection.
        
        Args:
            minio_endpoint: MinIO server endpoint
        """
        self.minio_endpoint = minio_endpoint
        
        if HAS_MINIO:
            try:
                self.minio_client = Minio(
                    minio_endpoint,
                    access_key=os.getenv("AWS_ACCESS_KEY_ID", "minioadmin"),
                    secret_key=os.getenv("AWS_SECRET_ACCESS_KEY", "minioadmin"),
                    secure=False
                )
                logger.info(" MinIO client initialized")
            except Exception as e:
                logger.error(f" Failed to initialize MinIO client: {e}")
                self.minio_client = None
        else:
            self.minio_client = None
    
    def extract_text(self, file_key: str, bucket: str = "bronze-resumes") -> str:
        """
        Extract text from PDF or DOCX file stored in MinIO.
        
        Args:
            file_key: Object key in MinIO bucket
            bucket: Bucket name (default: bronze-resumes)
            
        Returns:
            Extracted text content
        """
        try:
            if self.minio_client is None:
                # Fallback: read from local filesystem
                return self._extract_from_local(file_key)
            
            logger.info(f" Extracting text from: {file_key}")
            
            # Get file from MinIO
            response = self.minio_client.get_object(bucket, file_key)
            file_data = response.read()
            response.close()
            response.release_conn()
            
            # Determine file type and extract
            if file_key.lower().endswith('.pdf'):
                text = self._extract_from_pdf(file_data)
            elif file_key.lower().endswith('.docx'):
                text = self._extract_from_docx(file_data)
            else:
                raise ValueError(f"Unsupported file type: {file_key}")
            
            logger.info(f" Extracted {len(text)} characters from {file_key}")
            return text
            
        except Exception as e:
            logger.error(f" Error extracting text from {file_key}: {e}")
            return ""
    
    def _extract_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text from PDF bytes.
        Tries native extraction first, falls back to OCR if needed.
        
        Args:
            pdf_bytes: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        text = ""
        
        # Try native text extraction first
        if HAS_PDF:
            try:
                pdf_file = io.BytesIO(pdf_bytes)
                reader = PyPDF2.PdfReader(pdf_file)
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if len(text.strip()) > 100:
                    logger.info(" Native PDF text extraction successful")
                    return text
            except Exception as e:
                logger.warning(f" Native PDF extraction failed: {e}")
        
        # Fall back to OCR if native extraction failed
        if HAS_OCR and len(text.strip()) < 100:
            try:
                logger.info(" Attempting OCR extraction...")
                images = convert_from_bytes(pdf_bytes)
                
                ocr_text = ""
                for i, image in enumerate(images):
                    page_text = pytesseract.image_to_string(image)
                    ocr_text += page_text + "\n"
                    logger.info(f" OCR page {i+1}: {len(page_text)} chars")
                
                if len(ocr_text.strip()) > len(text.strip()):
                    logger.info(" OCR extraction successful")
                    return ocr_text
                    
            except Exception as e:
                logger.error(f" OCR extraction failed: {e}")
        
        return text if text else "Error: Could not extract text from PDF"
    
    def _extract_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.
        
        Args:
            docx_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        if not HAS_DOCX:
            return "Error: python-docx not installed"
        
        try:
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f" DOCX extraction successful: {len(text)} chars")
            return text
            
        except Exception as e:
            logger.error(f" DOCX extraction failed: {e}")
            return f"Error: {str(e)}"
    
    def _extract_from_local(self, file_path: str) -> str:
        """
        Extract text from local file (fallback when MinIO unavailable).
        
        Args:
            file_path: Path to local file
            
        Returns:
            Extracted text
        """
        try:
            with open(file_path, 'rb') as f:
                file_bytes = f.read()
            
            if file_path.lower().endswith('.pdf'):
                return self._extract_from_pdf(file_bytes)
            elif file_path.lower().endswith('.docx'):
                return self._extract_from_docx(file_bytes)
            else:
                return "Error: Unsupported file type"
                
        except Exception as e:
            logger.error(f" Local file extraction failed: {e}")
            return f"Error: {str(e)}"
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\-.,@()]', '', text)
        
        # Normalize to lowercase for consistency
        text = text.lower().strip()
        
        return text
    
    def extract_metadata(self, file_key: str) -> Dict[str, any]:
        """
        Extract metadata about the resume file.
        
        Args:
            file_key: Object key in MinIO
            
        Returns:
            Dictionary with metadata
        """
        metadata = {
            'file_key': file_key,
            'file_name': os.path.basename(file_key),
            'file_type': os.path.splitext(file_key)[1],
            'extracted_at': datetime.utcnow().isoformat(),
            'parser_version': '1.0.0'
        }
        
        # Get file size if MinIO available
        if self.minio_client:
            try:
                stat = self.minio_client.stat_object("bronze-resumes", file_key)
                metadata['file_size_bytes'] = stat.size
                metadata['upload_date'] = stat.last_modified.isoformat()
            except:
                pass
        
        return metadata


# Example usage
if __name__ == "__main__":
    parser = ResumeParser()
    
    # Test with sample file
    test_file = "sample_resume.pdf"
    text = parser.extract_text(test_file)
    print(f"Extracted text length: {len(text)}")
    print(f"First 200 chars: {text[:200]}")
